# ============================================================
# 长期记忆 — Chroma 向量库 (ChromaKnowledgeStore)
# ============================================================

from __future__ import annotations

import os
import logging
from pathlib import Path
from typing import Any

from langchain_chroma import Chroma
from langchain_core.documents import Document

from infra import resolve_reranker_base_url
from rag.embeddings import build_cloud_embeddings
from rag.reranker import ApiReranker, rerank_documents

logger = logging.getLogger(__name__)


class ChromaKnowledgeStore:
    """长期记忆 / 知识库：基于 LangChain Chroma 的语义向量存储。

    支持云端 Embedding + Cross-Encoder 远程重排。
    别名：LongTermMemory
    """

    def __init__(self, settings) -> None:
        self._settings = settings
        self._embeddings = build_cloud_embeddings(settings)
        self._persist_dir = settings.chroma_persist_dir
        self._reranker = ApiReranker(
            base_url=resolve_reranker_base_url(settings),
            api_key=settings.rag_reranker_api_key or settings.llm_api_key,
            model=settings.rag_reranker_model,
        )
        self._vectorstore: Chroma | None = None
        self._init_vectorstore()

    def _init_vectorstore(self) -> None:
        """初始化或加载 Chroma 向量库。"""
        os.makedirs(self._persist_dir, exist_ok=True)
        try:
            self._vectorstore = Chroma(
                persist_directory=self._persist_dir,
                embedding_function=self._embeddings,
                collection_name="knowledge_base",
            )
            count = self._vectorstore._collection.count() if self._vectorstore._collection else 0
            logger.info("ChromaKnowledgeStore initialized: %d documents in %s", count, self._persist_dir)
        except Exception as exc:
            logger.warning("Failed to initialize Chroma: %s — creating empty store", exc)
            self._vectorstore = None

    @property
    def is_ready(self) -> bool:
        return self._vectorstore is not None and (
            self._vectorstore._collection is not None and
            self._vectorstore._collection.count() > 0
        )

    @property
    def document_count(self) -> int:
        if self._vectorstore and self._vectorstore._collection:
            return self._vectorstore._collection.count()
        return 0

    # ── 写入 ──

    async def add_document(self, content: str, source: str = "",
                           metadata: dict[str, Any] | None = None) -> str:
        """添加单条文档到向量库。"""
        meta = metadata or {}
        meta["source"] = source
        doc = Document(page_content=content, metadata=meta)
        if self._vectorstore:
            ids = self._vectorstore.add_documents([doc])
            return ids[0] if ids else ""
        return ""

    async def load_knowledge_base(self, kb_dir: str) -> int:
        """批量加载 .txt 知识库文件到 Chroma（全量）。"""
        kb_path = Path(kb_dir)
        if not kb_path.exists():
            logger.warning("Knowledge base directory not found: %s", kb_dir)
            return 0

        txt_files = [str(f) for f in kb_path.rglob("*.txt")]
        if not txt_files:
            logger.warning("No .txt files found in %s", kb_dir)
            return 0

        return await self.load_files(txt_files, base_dir=kb_dir)

    async def load_files(self, file_paths: list[str], base_dir: str = "") -> int:
        """增量加载指定文件列表到 Chroma。支持 txt/md/pdf/docx/xlsx。

        Args:
            file_paths: 文件绝对路径列表
            base_dir: 基准目录，用于计算 source 元数据中的相对路径

        Returns:
            加载的文档总数
        """
        base_path = Path(base_dir) if base_dir else Path.cwd()
        chunk_size = self._settings.rag_chunk_size
        chunk_overlap = self._settings.rag_chunk_overlap
        total = 0

        for file_path in file_paths:
            file_obj = Path(file_path)
            if not file_obj.exists():
                logger.warning("File not found: %s", file_path)
                continue

            # 计算相对路径作为 source
            try:
                source = str(file_obj.relative_to(base_path)).replace("\\", "/")
            except ValueError:
                source = file_obj.name

            # 按文件类型解析文本内容
            try:
                text = self._extract_text(file_obj)
                if not text or not text.strip():
                    logger.warning("No text content extracted from %s", source)
                    continue
            except Exception as exc:
                logger.warning("Failed to extract text from %s: %s", file_obj.name, exc)
                continue

            try:
                # 按空行分块
                chunks = [c.strip() for c in text.split("\n\n") if c.strip()]
                if not chunks:
                    chunks = [text.strip()]

                docs = []
                for i, chunk in enumerate(chunks):
                    if len(chunk) > chunk_size:
                        for j in range(0, len(chunk), chunk_size - chunk_overlap):
                            sub = chunk[j:j + chunk_size]
                            if sub.strip():
                                docs.append(Document(
                                    page_content=sub,
                                    metadata={"source": source, "chunk": i}
                                ))
                    else:
                        docs.append(Document(
                            page_content=chunk,
                            metadata={"source": source, "chunk": i}
                        ))

                if docs and self._vectorstore:
                    self._vectorstore.add_documents(docs)
                    total += len(docs)
                    logger.info("Loaded %d chunks from %s", len(docs), source)
            except Exception as exc:
                logger.warning("Failed to index %s: %s", file_obj.name, exc)

        logger.info("load_files done: %d total chunks from %d files", total, len(file_paths))
        return total

    @staticmethod
    def _extract_text(file_obj: Path) -> str:
        """根据文件扩展名提取文本内容。

        支持: .txt .md（纯文本）、.pdf（pymupdf）、.docx（python-docx）、.xlsx（openpyxl）
        """
        ext = file_obj.suffix.lower()

        # 纯文本
        if ext in (".txt", ".md"):
            return file_obj.read_text(encoding="utf-8", errors="replace")

        # PDF
        if ext == ".pdf":
            try:
                import fitz  # pymupdf
                text_parts: list[str] = []
                doc = fitz.open(str(file_obj))
                for page in doc:
                    page_text = page.get_text()
                    if page_text.strip():
                        text_parts.append(page_text.strip())
                doc.close()
                return "\n\n".join(text_parts)
            except ImportError:
                logger.warning("pymupdf not installed, cannot parse PDF: %s", file_obj.name)
                return ""

        # Word
        if ext == ".docx":
            try:
                from docx import Document as DocxDocument
                doc = DocxDocument(str(file_obj))
                paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
                return "\n\n".join(paragraphs)
            except ImportError:
                logger.warning("python-docx not installed, cannot parse docx: %s", file_obj.name)
                return ""

        # Excel
        if ext == ".xlsx":
            try:
                import openpyxl
                wb = openpyxl.load_workbook(str(file_obj), read_only=True, data_only=True)
                all_text: list[str] = []
                for sheet_name in wb.sheetnames:
                    ws = wb[sheet_name]
                    rows_text: list[str] = [f"## 工作表: {sheet_name}"]
                    for row in ws.iter_rows(values_only=True):
                        row_vals = [str(c) for c in row if c is not None]
                        if row_vals:
                            rows_text.append(" | ".join(row_vals))
                    if len(rows_text) > 1:
                        all_text.append("\n".join(rows_text))
                wb.close()
                return "\n\n".join(all_text)
            except ImportError:
                logger.warning("openpyxl not installed, cannot parse xlsx: %s", file_obj.name)
                return ""

        logger.warning("Unsupported file type: %s (%s)", ext, file_obj.name)
        return ""

    # ── 检索 ──

    async def search(self, query: str, top_k: int = 12) -> list[Document]:
        """纯向量相似度检索。"""
        if not self._vectorstore:
            return []
        try:
            return self._vectorstore.similarity_search(query, k=top_k)
        except Exception as exc:
            logger.warning("Vector search failed: %s", exc)
            return []

    async def search_and_rerank(self, query: str, top_k: int = 12,
                                 top_n: int = 5) -> list[Document]:
        """向量召回 + Cross-Encoder 重排。"""
        candidates = await self.search(query, top_k=top_k)
        if not candidates or len(candidates) <= top_n:
            return candidates
        return await rerank_documents(self._reranker, query, candidates, top_n=top_n)
